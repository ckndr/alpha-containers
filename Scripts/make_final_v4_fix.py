import docx
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import zipfile
import os

def set_cell_background(cell, color_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

def generate_v4():
    # Extract Logo
    docx_path = r'd:\Alpha\Aerosol\SOPs\AER-PL-001_Order_Receipt_and_Review.docx'
    temp_img = r'd:\Alpha\Aerosol\temp_logo.jpg'
    
    with zipfile.ZipFile(docx_path, 'r') as z:
        with open(temp_img, 'wb') as f:
            f.write(z.read('word/media/image1.jpg'))

    # Open Claude's original template
    doc = docx.Document(r"D:\Alpha\Aerosol\Aerosol_Job_Order_Form_JobCard.docx")

    # 1. Remove Word Document Header (The red highlighted part)
    for section in doc.sections:
        for t in section.header.tables:
            t._element.getparent().remove(t._element)
        for p in section.header.paragraphs:
            p.text = ""

    # 2. Insert new Logo Header at the top
    first_p = doc.paragraphs[0]
    table_hdr = doc.add_table(rows=1, cols=3)
    table_hdr.autofit = False
    first_p._element.addprevious(table_hdr._element)
    
    table_hdr.columns[0].width = Inches(1.5)
    table_hdr.columns[1].width = Inches(4.0)
    table_hdr.columns[2].width = Inches(2.5)
    
    # Logo
    cell_0 = table_hdr.cell(0, 0)
    p0 = cell_0.paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run0 = p0.add_run()
    run0.add_picture(temp_img, width=Inches(1.2))
    
    # Title
    cell_1 = table_hdr.cell(0, 1)
    p1 = cell_1.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run1 = p1.add_run("Alpha Aerosols\nJob Order Card")
    run1.font.bold = True
    run1.font.size = Pt(16)
    
    # Doc Info
    cell_2 = table_hdr.cell(0, 2)
    p2 = cell_2.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run2 = p2.add_run("DOC #: AER-JOC-001\nDATE: [Placeholder]\nREVISION: 00")
    run2.font.bold = True
    run2.font.size = Pt(10)

    # 3. Update Table 0 (Job Details)
    table_0 = doc.tables[0]
    # Remove all existing rows in table_0
    pass #
        tr = table_0.rows[0]._tr
        table_0._tbl.remove(tr)
        
    job_details = [
        ("Job Order No.:", "", "Date:", ""),
        ("Product / Can Size:", "", "Order Quantity (Nos.):", ""),
        ("Customer:", "", "Shift:", ""),
        ("Artwork / Design Ref.:", "", "", "")
    ]
    for row_data in job_details:
        row_cells = table_0.add_row().cells
        for i in range(4):
            row_cells[i].text = row_data[i]
        for col_idx in [0, 2]:
            set_cell_background(row_cells[col_idx], "DCE6F1")
            if row_cells[col_idx].text:
                for p in row_cells[col_idx].paragraphs:
                    for run in p.runs:
                        run.font.bold = True

    # 4. Update Table 1 (BOM)
    table_1 = doc.tables[1]
    hdr_cells = table_1.rows[0].cells
    for cell in hdr_cells:
        if "Qty per Unit" in cell.text:
            cell.text = "Total Required Qty"
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.bold = True
        elif "Tolerance" in cell.text:
            cell.text = "Waste / Scrap %"
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.bold = True

    # 5. Update Table 2 (Product Specifications)
    table_2 = doc.tables[2]
    # Clear Table 2
    for _ in range(len(table_2.rows)):
        tr = table_2.rows[0]._tr
        table_2._tbl.remove(tr)
        
    specs = [
        ("Can Type:", "", "Can Diameter (mm):", ""),
        ("Can Height (mm):", "", "Neck Opening (mm):", ""),
        ("Wall Thickness (mm):", "", "Dome/Cone Thickness:", ""),
        ("Internal Coating:", "", "Base Coat Color:", ""),
        ("External Print Colors:", "", "Varnish Type:", ""),
        ("Qty per Pallet/Carton:", "", "Special Instructions:", "")
    ]
    for row_data in specs:
        row_cells = table_2.add_row().cells
        for i in range(4):
            row_cells[i].text = row_data[i]
        for col_idx in [0, 2]:
            set_cell_background(row_cells[col_idx], "DCE6F1")
            if row_cells[col_idx].text:
                for p in row_cells[col_idx].paragraphs:
                    for run in p.runs:
                        run.font.bold = True

    # Save
    out_path = r"D:\Alpha\Aerosol\Aerosol_Job_Card_Final_V4.docx"
    doc.save(out_path)
    print(f"Generated {out_path}")
    
    if os.path.exists(temp_img):
        os.remove(temp_img)

if __name__ == "__main__":
    generate_v4()
